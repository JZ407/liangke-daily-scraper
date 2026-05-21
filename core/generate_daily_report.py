"""
量科网每日新闻日报自动生成器
用法: python generate_daily_report.py [日期，如 2026.5.20]
不填日期则默认今天
"""
import sys
import re
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from db import get_session, Article


# 自动标签关键词库（用于文章重要性评分）
SCORE_KEYWORDS = {
    'commercial': ['融资', '估值', '亿美元', '万美元', '营收', 'ipo', '上市', '收购', '合并', '推出', '发布', '启用'],
    'partnership': ['合作', '签署', '达成', '战略', '联盟', '协议'],
    'policy': ['nist', '标准化', '标准', '政策', '计划', '路线图'],
    'major_companies': ['ibm', 'google', '谷歌', '富士通', 'pasqal', 'ionq', 'quantinuum', 'infleqtion', '英伟达', 'nvidia', '微软', 'microsoft', '亚马逊', 'amazon'],
}


def score_article(a: Article) -> int:
    """计算文章重要性得分。"""
    s = 0
    text = (a.title or '') + ' ' + (a.content or '')
    text_lower = text.lower()

    # 新闻类型优先
    if '/flash/' in a.liangke_url or '/article/' in a.liangke_url:
        s += 10
    else:
        s -= 5  # 学术文章降权

    # 商业/融资关键词
    for kw in SCORE_KEYWORDS['commercial']:
        if kw in text_lower:
            s += 3

    # 合作关键词
    for kw in SCORE_KEYWORDS['partnership']:
        if kw in text_lower:
            s += 2

    # 政策/标准
    for kw in SCORE_KEYWORDS['policy']:
        if kw in text_lower:
            s += 2

    # 知名企业
    for c in SCORE_KEYWORDS['major_companies']:
        if c in text_lower:
            s += 2

    # 时效性
    if a.original_date and a.original_date >= date.today() - timedelta(days=1):
        s += 2

    return s


def clean_content(content: str) -> str:
    """删除内容开头的日期前缀，避免重复。"""
    if not content:
        return content
    patterns = [
        r'^\d{4}年\d{1,2}月\d{1,2}日\s*[\-–—:：,，\s]*',
        r'^\d{1,2}月\d{1,2}日\s*[\-–—:：,，\s]*',
        r'^\d{4}-\d{2}-\d{2}\s*[\-–—:：,，\s]*',
    ]
    for pat in patterns:
        content, count = re.subn(pat, '', content, count=1)
        if count > 0:
            break
    return content.strip()


def generate_report(target_date: date = None) -> str:
    """生成日报 docx，返回文件路径。"""
    if target_date is None:
        target_date = date.today()

    session = get_session()

    # 取最近 2 天的文章作为候选池
    recent = session.query(Article).filter(
        Article.liangke_date >= target_date - timedelta(days=2)
    ).all()

    if len(recent) < 3:
        print(f"WARNING: Only {len(recent)} articles found, need at least 3.")

    # 评分并取 top3
    scored = [(a, score_article(a)) for a in recent]
    scored.sort(key=lambda x: x[1], reverse=True)
    top3 = [a for a, _ in scored[:3]]

    # 创建文档
    doc = Document()

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f'今日快讯（{target_date.year}-{target_date.month}-{target_date.day}）：')
    title_run.font.name = 'Microsoft YaHei'
    title_run.font.size = Pt(18)   # 小二
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # 新闻正文
    for i, a in enumerate(top3, 1):
        # 新闻标题（三号）
        t_para = doc.add_paragraph()
        t_run = t_para.add_run(f'{i}、{a.title}')
        t_run.font.name = 'Microsoft YaHei'
        t_run.font.size = Pt(16)   # 三号
        t_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 正文（四号）
        cleaned = clean_content(a.content or '')
        orig = a.original_date or target_date
        date_str = f'{orig.month}月{orig.day}日'
        content_text = f'{date_str}消息，{cleaned}'

        # 按自然段拆分，每段单独设置首行缩进
        paragraphs = content_text.split('\n')
        for idx, para_text in enumerate(paragraphs):
            if not para_text.strip():
                continue
            c_para = doc.add_paragraph()
            c_para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
            c_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 靠左对齐
            c_run = c_para.add_run(para_text)
            c_run.font.name = 'Microsoft YaHei'
            c_run.font.size = Pt(14)   # 四号
            c_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 链接标签（五号）
        l_label_para = doc.add_paragraph()
        l_label_run = l_label_para.add_run('链接：')
        l_label_run.font.name = 'Microsoft YaHei'
        l_label_run.font.size = Pt(10.5)   # 五号
        l_label_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # URL（五号）
        url = a.reference_url if a.reference_url and a.reference_url != a.liangke_url else a.liangke_url
        u_para = doc.add_paragraph()
        u_run = u_para.add_run(url)
        u_run.font.name = 'Microsoft YaHei'
        u_run.font.size = Pt(10.5)   # 五号
        u_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        doc.add_paragraph()

    # 专利占位
    doc.add_paragraph()
    p_para = doc.add_paragraph()
    p_run = p_para.add_run('4、专利：（请自行填写）')
    p_run.font.name = 'Microsoft YaHei'
    p_run.font.size = Pt(16)   # 三号
    p_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()   # 用户填写内容的位置

    p_link_label = doc.add_paragraph()
    p_link_run = p_link_label.add_run('链接：')
    p_link_run.font.name = 'Microsoft YaHei'
    p_link_run.font.size = Pt(10.5)   # 五号
    p_link_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()   # 用户填写 URL 的位置

    filename = f'日报{target_date.year}.{target_date.month}.{target_date.day}.docx'
    doc.save(filename)
    print(f'日报已生成: {filename}')
    print('入选新闻:')
    for i, a in enumerate(top3, 1):
        print(f'  {i}. [{a.id}] {a.title}')
    return filename


def main():
    if len(sys.argv) > 1:
        # 支持格式: 2026.5.20 或 2026-05-20
        date_str = sys.argv[1].replace('-', '.')
        parts = date_str.split('.')
        target = date(int(parts[0]), int(parts[1]), int(parts[2]))
    else:
        target = date.today()

    generate_report(target)


if __name__ == '__main__':
    main()
